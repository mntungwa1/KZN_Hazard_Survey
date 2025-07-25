import streamlit as st
import pandas as pd
import geopandas as gpd
from shapely.geometry import Point
from streamlit_folium import st_folium
import folium
from pathlib import Path
from datetime import datetime, date
import os, smtplib, re, zipfile
from email.message import EmailMessage
from docx import Document
from fpdf import FPDF

# --- Configuration ---
BASE_DIR = Path("C:/Temp/kzn")
SAVE_DIR = BASE_DIR / "Responses"
SAVE_DIR.mkdir(parents=True, exist_ok=True)
MASTER_CSV = BASE_DIR / "all_submissions.csv"
EXCEL_PATH = Path("RiskAssessmentTool.xlsm")
GEOJSON_PATH = Path("KZN_wards.geojson")

# --- Load from Streamlit Secrets ---
EMAIL_ADDRESS = st.secrets["EMAIL_ADDRESS"]
EMAIL_PASSWORD = st.secrets["EMAIL_PASSWORD"]
APP_PASSWORD = st.secrets.get("APP_PASSWORD", "kzn!23@")
ADMIN_PASSWORD = st.secrets.get("ADMIN_PASSWORD", "kzn!23&")
ADMIN_EMAILS = [st.secrets.get("ADMIN_EMAIL", EMAIL_ADDRESS), "dingaanm@gmail.com"]

LOGO_PATH = "Logo.png"
SRK_LOGO_PATH = "SRK_Logo.png"

# --- Setup ---
def ensure_save_dir():
    SAVE_DIR.mkdir(parents=True, exist_ok=True)
    MASTER_CSV.parent.mkdir(parents=True, exist_ok=True)

ensure_save_dir()

def safe_filename(name):
    return re.sub(r'[^A-Za-z0-9_-]', '_', name)

# --- Authentication ---
def password_protection():
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False
    password = st.text_input("Enter password to access the app:", type="password")
    if st.button("Login"):
        if password == APP_PASSWORD:
            st.session_state["authenticated"] = True
            st.success("Access granted. Please continue.")
            st.rerun()
        else:
            st.error("Incorrect password.")

if not st.session_state.get("authenticated", False):
    st.title("KZN Hazard Risk Assessment Survey - Login")
    password_protection()
    st.stop()

# --- Load Hazards ---
@st.cache_data(show_spinner=False, ttl=3600)
def load_hazards():
    df = pd.read_excel(EXCEL_PATH, sheet_name="Hazard information", skiprows=1)
    return df.iloc[:, 0].dropna().tolist()

# --- Load Wards ---
@st.cache_data(show_spinner=False, ttl=3600)
def load_ward_gdf():
    return gpd.read_file(GEOJSON_PATH).to_crs(epsg=4326)

# --- Email Sending ---
def send_email(subject, body, to_emails, attachments):
    try:
        msg = EmailMessage()
        msg["Subject"] = subject
        msg["From"] = EMAIL_ADDRESS
        msg["To"] = ", ".join(to_emails)
        msg.set_content(body)
        for attachment in attachments:
            with open(attachment, "rb") as f:
                file_data = f.read()
                msg.add_attachment(file_data, maintype="application", subtype="octet-stream", filename=Path(attachment).name)
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            smtp.send_message(msg)
        st.success(f"Email sent to {to_emails}!")
    except Exception as e:
        st.error(f"Failed to send email: {e}")

# --- Save Responses ---
def append_to_master_csv(df):
    df.to_csv(MASTER_CSV, mode="a", header=not MASTER_CSV.exists(), index=False)

def save_responses(responses, name, ward, email, date_filled,
                   district_municipality=None, local_municipality=None, extra_info=None):
    ensure_save_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_filename = f"{safe_filename(ward)}_{safe_filename(name)}_{timestamp}"
    csv_path = SAVE_DIR / f"{base_filename}.csv"
    pdf_path = SAVE_DIR / f"{base_filename}.pdf"
    docx_path = SAVE_DIR / f"{base_filename}.docx"

    df = pd.DataFrame(responses)
    df.insert(0, "Respondent Name", name)
    df.insert(1, "District Municipality", district_municipality)
    df.insert(2, "Local Municipality", local_municipality)
    df.insert(3, "Ward", ward)
    df.insert(4, "Email", email)
    df.insert(5, "Extra Info", extra_info)
    df.insert(6, "Date", date_filled)
    df.to_csv(csv_path, index=False)
    append_to_master_csv(df)

    # DOCX
    doc = Document()
    doc.add_heading("KZN Hazard Risk Assessment Survey", 0)
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"District Municipality: {district_municipality}")
    doc.add_paragraph(f"Local Municipality: {local_municipality}")
    doc.add_paragraph(f"Ward: {ward}")
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Extra Info: {extra_info}")
    doc.add_paragraph(f"Date: {date_filled}")
    doc.add_paragraph("---")
    for _, row in df.iterrows():
        doc.add_paragraph(f"Hazard: {row['Hazard']} | Question: {row['Question']} | Response: {row['Response']}")
    doc.save(docx_path)

    # PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(200, 10, txt="KZN Hazard Risk Assessment Survey", ln=True, align="C")
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Name: {name}", ln=True)
    pdf.cell(200, 10, txt=f"District Municipality: {district_municipality}", ln=True)
    pdf.cell(200, 10, txt=f"Local Municipality: {local_municipality}", ln=True)
    pdf.cell(200, 10, txt=f"Ward: {ward}", ln=True)
    pdf.cell(200, 10, txt=f"Email: {email}", ln=True)
    pdf.multi_cell(0, 10, txt=f"Extra Info: {extra_info}")
    pdf.cell(200, 10, txt=f"Date: {date_filled}", ln=True)
    pdf.ln(5)
    for _, row in df.iterrows():
        pdf.multi_cell(0, 10, txt=f"Hazard: {row['Hazard']} | Question: {row['Question']} | Response: {row['Response']}")
    pdf.output(pdf_path)

    return csv_path, docx_path, pdf_path

# --- Create ZIP ---
def create_zip(local_municipality, files):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_name = f"{safe_filename(local_municipality)}_{timestamp}.zip"
    zip_path = SAVE_DIR / zip_name
    with zipfile.ZipFile(zip_path, "w") as zipf:
        for file in files:
            zipf.write(file, os.path.basename(file))
    return zip_path

# --- Hazard Questions ---
questions_with_descriptions = {
    "Has this hazard occurred in the past?": [
        "0 - Has not occurred and has no chance of occurrence",
        "1 - Has not occurred but there is real potential for occurrence",
        "2 - Has occurred but only once",
        "3 - Has occurred but only a few times or rarely",
        "4 - Has occurred regularly or at least once a year",
        "5 - Occurs multiple times during a single year",
    ],
    "How frequently does it occur?": [
        "0 - Unknown / Not applicable",
        "1 - Decreasing",
        "2 - Stable",
        "3 - Marginally increasing",
        "4 - Increasing",
        "5 - Increasing rapidly",
    ],
    "What is the typical duration of the hazard?": [
        "0 - Unknown / Not applicable",
        "1 - Few minutes",
        "2 - Few hours",
        "3 - Few days",
        "4 - Few weeks",
        "5 - Few months",
    ],
    "What is the area of impact?": [
        "0 - None",
        "1 - Single property",
        "2 - Single Ward",
        "3 - Few wards",
        "4 - Entire municipality",
        "5 - Larger than municipality",
    ],
    "What is the impact on people?": [
        "0 - None",
        "1 - Low impact / Discomfort",
        "2 - Minimal impact / Minor injuries",
        "3 - Serious injuries / Health problems no fatalities",
        "4 - Fatalities / Serious health problems but confined",
        "5 - Multiple fatalities spread over wide area",
    ],
    "What is the impact on infrastructure and services?": [
        "0 - None",
        "1 - Low impact / Minor damage / Minor disruption",
        "2 - Some structural damage / Short term disruption of services",
        "3 - Medium structural damage / 1 Week disruption",
        "4 - Serious structural damage / Disruption of longer than a week",
        "5 - Total disruption of structure / Disruption of longer than a month",
    ],
    "What is the impact on the environment?": [
        "0 - Not applicable / No effects",
        "1 - Minor effects",
        "2 - Medium effects",
        "3 - Severe",
        "4 - Severe effects over wide area",
        "5 - Total destruction",
    ],
    "What is the level of economic disruption?": [
        "0 - No disruption",
        "1 - Some disruption",
        "2 - Medium disruption",
        "3 - Severe short-term disruption",
        "4 - Severe long-term disruption",
        "5 - Total stop in activities",
    ],
    "How predictable is the hazard?": [
        "0 - Not applicable",
        "1 - Effective early warning",
        "3 - Partially predictable",
        "5 - No early warning",
    ],
    "What is the urgency or priority level?": [
        "0 - Not applicable / No effects",
        "1 - Low priority",
        "2 - Medium priority",
        "3 - Medium high priority",
        "4 - High priority",
        "5 - Very high priority",
    ],
}

capacity_questions = [
    "Sufficient staff/human resources",
    "Experience and special knowledge",
    "Equipment availability",
    "Adequate funding/budget allocation",
    "Facilities and infrastructure for response",
    "Prevention and mitigation plans",
    "Response and recovery plans",
    "Community awareness and training programs",
    "Early warning systems in place",
    "Coordination with local authorities and partners",
]

capacity_options = [
    "Strongly Disagree",
    "Disagree",
    "Neutral",
    "Agree",
    "Strongly Agree",
]

def build_hazard_questions(hazards_to_ask):
    responses = []
    for hazard in hazards_to_ask:
        st.markdown(f"### {hazard}")
        for q, opts in questions_with_descriptions.items():
            response = st.radio(q, opts, key=f"{hazard}_{q}")
            responses.append({"Hazard": hazard, "Question": q, "Response": response})
        for cq in capacity_questions:
            response = st.radio(cq, capacity_options, key=f"{hazard}_{cq}")
            responses.append({"Hazard": hazard, "Question": cq, "Response": response})
    return responses

# --- Map Display with Hover Highlight ---
def display_map(gdf):
    m = folium.Map(location=[-29.5, 31.1], zoom_start=7)
    folium.GeoJson(
        data=gdf.__geo_interface__,
        style_function=lambda x: {"fillColor": "#3186cc", "color": "black", "weight": 1, "fillOpacity": 0.4},
        highlight_function=lambda x: {"fillColor": "#ffcc00", "color": "black", "weight": 2, "fillOpacity": 0.7},
        tooltip=folium.GeoJsonTooltip(fields=[gdf.columns[0]], aliases=["UID:"], sticky=True)
    ).add_to(m)
    return st_folium(m, height=1000, width=1200)

# --- Survey ---
def run_survey():
    st.title("KZN Hazard Risk Assessment Survey")
    hazards = load_hazards()
    gdf = load_ward_gdf()
    map_data = display_map(gdf)

    clicked_ward = None
    if map_data.get("last_clicked"):
        pt = Point(map_data["last_clicked"]["lng"], map_data["last_clicked"]["lat"])
        for _, row in gdf.iterrows():
            if row.geometry.contains(pt):
                clicked_ward = row[gdf.columns[0]]
                st.session_state["selected_ward"] = clicked_ward
                break

    ward_display = st.session_state.get("selected_ward", "")
    if ward_display:
        st.success(f"Selected Ward: {ward_display}")

    st.subheader("Select Applicable Hazards")
    selected = st.multiselect("Choose hazards:", hazards)
    custom = st.text_input("Other hazard") if st.checkbox("Add custom hazard") else ""

    if selected or custom:
        if "active_tab" not in st.session_state:
            st.session_state.active_tab = "Respondent Info"

        if st.session_state.active_tab == "Respondent Info":
            st.subheader("Respondent Info")
            st.session_state["name"] = st.text_input("Full Name", st.session_state.get("name", ""))
            st.session_state["district_municipality"] = st.text_input("District Municipality", st.session_state.get("district_municipality", ""))
            st.session_state["local_municipality"] = st.text_input("Local Municipality", st.session_state.get("local_municipality", ""))
            st.session_state["final_ward"] = ward_display or st.text_input("Ward (if not using map)", st.session_state.get("final_ward", ""))
            st.session_state["today"] = st.date_input("Date", value=st.session_state.get("today", date.today()))
            st.session_state["user_email"] = st.text_input("Your Email", st.session_state.get("user_email", ""))
            st.session_state["extra_info"] = st.text_area("Any extra information to be added", st.session_state.get("extra_info", ""))

            if st.button("Click Hazard Risk Evaluation Tab"):
                st.session_state.active_tab = "Hazard Risk Evaluation"
                st.rerun()

        elif st.session_state.active_tab == "Hazard Risk Evaluation":
            st.subheader("Hazard Risk Evaluation")
            hazards_to_ask = selected + ([custom] if custom else [])
            with st.form("hazard_form"):
                responses = build_hazard_questions(hazards_to_ask)
                col1, col2 = st.columns(2)
                with col1:
                    back = st.form_submit_button("Go Back to Respondent Info Tab")
                with col2:
                    submit = st.form_submit_button("Submit Survey")
                if back:
                    st.session_state.active_tab = "Respondent Info"
                    st.rerun()
                if submit:
                    if not st.session_state.get("name") or not st.session_state.get("final_ward"):
                        st.error("Please fill in your name and ward.")
                    else:
                        csv_file, doc_file, pdf_file = save_responses(
                            responses,
                            st.session_state["name"],
                            st.session_state["final_ward"],
                            st.session_state["user_email"],
                            st.session_state["today"],
                            st.session_state["district_municipality"],
                            st.session_state["local_municipality"],
                            st.session_state["extra_info"]
                        )
                        zip_file = create_zip(st.session_state["local_municipality"], [csv_file, doc_file, pdf_file])
                        st.session_state["files_saved"] = (csv_file, doc_file, pdf_file, zip_file)
                        st.success(f"Survey submitted successfully! Files saved in: {SAVE_DIR}")

                        # Send ZIP via Email
                        if st.session_state["user_email"]:
                            send_email(
                                "Your KZN Hazard Survey Submission",
                                "Thank you for completing the survey. Your files are attached as a ZIP archive.",
                                [st.session_state["user_email"]],
                                [zip_file]
                            )
                        send_email(
                            "New KZN Hazard Survey Submission",
                            "A new survey has been submitted. See attached ZIP file.",
                            ADMIN_EMAILS,
                            [zip_file]
                        )

        # Show download buttons outside the form
        if "files_saved" in st.session_state:
            csv_file, doc_file, pdf_file, zip_file = st.session_state["files_saved"]
            with open(csv_file, "rb") as f:
                st.download_button("Download CSV", f, file_name=os.path.basename(csv_file), mime="text/csv")
            with open(doc_file, "rb") as f:
                st.download_button("Download DOCX", f, file_name=os.path.basename(doc_file),
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with open(pdf_file, "rb") as f:
                st.download_button("Download PDF", f, file_name=os.path.basename(pdf_file), mime="application/pdf")
            with open(zip_file, "rb") as zf:
                st.download_button("Download All (ZIP)", zf, file_name=os.path.basename(zip_file), mime="application/zip")

# --- Main App ---
st.set_page_config(page_title="KZN Hazard Risk Assessment", layout="wide")
st.markdown("<style>div.block-container{padding-top: 1rem;}</style>", unsafe_allow_html=True)

menu = st.sidebar.radio("Navigation", ["Survey", "Admin Dashboard"])

if os.path.exists(LOGO_PATH): st.sidebar.image(LOGO_PATH, width=180)
if os.path.exists(SRK_LOGO_PATH): st.sidebar.image(SRK_LOGO_PATH, width=160)

# Disclaimer
st.sidebar.markdown(
    "<small><i>Disclaimer: The software is developed by Dingaan Mahlangu and should not be used without prior permission.</i></small>",
    unsafe_allow_html=True
)

if menu == "Survey":
    run_survey()
elif menu == "Admin Dashboard":
    st.title("Admin Dashboard - KZN Hazard Survey")
    if "admin_authenticated" not in st.session_state:
        st.session_state["admin_authenticated"] = False
    if not st.session_state["admin_authenticated"]:
        admin_password = st.text_input("Enter Admin Password:", type="password")
        if st.button("Login as Admin"):
            if admin_password == ADMIN_PASSWORD:
                st.session_state["admin_authenticated"] = True
                st.success("Admin Access Granted.")
                st.rerun()
            else:
                st.error("Incorrect Admin Password.")
        st.stop()
    if MASTER_CSV.exists():
        df = pd.read_csv(MASTER_CSV)
        st.dataframe(df)
        st.download_button("Download CSV", df.to_csv(index=False).encode("utf-8"),
                           file_name="filtered_submissions.csv", mime="text/csv")
    else:
        st.warning("No submissions found.")

# -*- coding: utf-8 -*-
import streamlit as st
import pandas as pd
import geopandas as gpd
import os, shutil, zipfile, re, smtplib, sys
from datetime import datetime, date
from shapely.geometry import Point
from streamlit_folium import st_folium
from email.message import EmailMessage
from dotenv import load_dotenv
import folium
from docx import Document
from fpdf import FPDF
from pathlib import Path
import webbrowser

# --- Load environment variables ---
load_dotenv()
EMAIL_ADDRESS = os.getenv("EMAIL_ADDRESS", "dummy_email@gmail.com")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD", "dummy_password")
APP_PASSWORD = os.getenv("APP_PASSWORD", "kzn!23@")

# -------------------------------
# PASSWORD PROTECTION SECTION
# -------------------------------
def password_protection():
    password = st.text_input("Enter password to access the app:", type="password")
    if password == APP_PASSWORD:
        st.session_state["authenticated"] = True
        st.success("Access granted.")
    elif password:
        st.error("Incorrect password.")

if "authenticated" not in st.session_state or not st.session_state["authenticated"]:
    st.title("🔒 KZN Hazard Risk Assessment Survey - Login")
    password_protection()
    st.stop()
# -------------------------------

# --- File paths ---
BASE_DIR = Path("C:/tmp/kzn")
TODAY_FOLDER = datetime.now().strftime("%d_%b_%Y")
SAVE_DIR = BASE_DIR / TODAY_FOLDER
SAVE_DIR.mkdir(parents=True, exist_ok=True)

EXCEL_PATH = Path("RiskAssessmentTool.xlsm")
GEOJSON_PATH = Path("KZN_wards.geojson")
LOGO_PATH = "Logo.png"
SRK_LOGO_PATH = "SRK_Logo.png"

# --- Cleanup old folders ---
def cleanup_old_folders(base_dir, days=30):
    now = datetime.now()
    pattern = re.compile(r"\d{2}_[A-Za-z]{3}_\d{4}")
    for folder in base_dir.iterdir():
        if folder.is_dir() and pattern.fullmatch(folder.name):
            try:
                if (now - datetime.strptime(folder.name, "%d_%b_%Y")).days > days:
                    shutil.rmtree(folder)
            except:
                continue
cleanup_old_folders(BASE_DIR)

# --- Load data ---
@st.cache_data
def load_hazards():
    try:
        df = pd.read_excel(EXCEL_PATH, sheet_name="Hazard information", skiprows=1)
        hazards_list = df.iloc[:, 0].dropna().tolist()
        if not hazards_list:
            st.error("No hazards found in the 'Hazard information' sheet. Please check the Excel file.")
            st.stop()
        return hazards_list
    except Exception as e:
        st.error("Error reading hazards: {}".format(e))
        st.stop()

@st.cache_data
def load_ward_gdf():
    return gpd.read_file(GEOJSON_PATH).to_crs(epsg=4326)

hazards = load_hazards()
gdf = load_ward_gdf()

# --- Page layout ---
st.set_page_config(page_title="KZN Hazard Risk Assessment", layout="wide")
st.image(LOGO_PATH, width=240)
st.image(SRK_LOGO_PATH, width=200)
st.title("KZN Hazard Risk Assessment Survey")
st.markdown("---")

# --- Interactive Map ---
st.subheader("Select a Ward from the Map")
m = folium.Map(location=[-29.5, 31.1], zoom_start=7)
folium.GeoJson(
    data=gdf.__geo_interface__,
    name="Wards",
    tooltip=folium.GeoJsonTooltip(fields=gdf.columns[:1].tolist()),
    popup=folium.GeoJsonPopup(fields=gdf.columns[:1].tolist()),
    highlight_function=lambda x: {"fillColor": "#ffaf00", "color": "black", "weight": 2},
).add_to(m)

map_data = st_folium(m, height=500)

# --- Ward selection logic ---
clicked_ward = None
if "last_active_drawing" in map_data and map_data["last_active_drawing"]:
    props = map_data["last_active_drawing"].get("properties", {})
    clicked_ward = props.get(gdf.columns[0])
elif "last_clicked" in map_data and map_data["last_clicked"]:
    lng = map_data["last_clicked"]["lng"]
    lat = map_data["last_clicked"]["lat"]
    pt = Point(lng, lat)
    for _, row in gdf.iterrows():
        if row.geometry.contains(pt):
            clicked_ward = row[gdf.columns[0]]
            break

if clicked_ward:
    st.session_state["selected_ward"] = clicked_ward
ward_display = st.session_state.get("selected_ward", "")
if ward_display:
    st.success("Selected Ward: {}".format(ward_display))

# --- Hazard Selection ---
st.markdown("---")
st.subheader("Select Applicable Hazards")
selected = st.multiselect("Choose hazards:", hazards)
include_other = st.checkbox("Add a custom hazard (Other)")
custom = st.text_input("Specify other hazard:") if include_other else ""

# --- Survey Form ---
submitted = False
if selected or custom:
    with st.form("hazard_form"):
        tab1, tab2 = st.tabs(["Respondent Info", "Hazard Risk Evaluation"])

        with tab1:
            name = st.text_input("Full Name")
            final_ward = ward_display or st.text_input("Ward (if not using map)")
            today = st.date_input("Date", value=date.today())
            user_email = st.text_input("Your Email")
            confirm = st.checkbox("I confirm the information is accurate")

        with tab2:
            levels = ["0 - Not applicable", "1 - Low", "2 - Moderate", "3 - High", "4 - Severe"]
            score_map = {v: i for i, v in enumerate(levels)}
            hazards_to_ask = selected + ([custom] if custom else [])
            responses = []
            all_filled = False  # Track if at least one hazard entry is filled
            for hazard in hazards_to_ask:
                st.markdown("**{}**".format(hazard))
                like = st.selectbox("Likelihood:", levels, key="{}_like".format(hazard))
                impact = st.selectbox("Impact:", levels, key="{}_impact".format(hazard))
                disrupt = st.selectbox("Disruption:", levels, key="{}_disrupt".format(hazard))
                if like != levels[0] or impact != levels[0] or disrupt != levels[0]:
                    all_filled = True
                risk_score = score_map[like] * score_map[impact] * score_map[disrupt]
                responses.append({
                    "Name": name,
                    "Ward": final_ward,
                    "Date": today,
                    "Hazard": hazard,
                    "Likelihood": like,
                    "Impact": impact,
                    "Disruption": disrupt,
                    "Risk Score": risk_score
                })

        submitted = st.form_submit_button("Submit Survey", disabled=not all_filled)

# --- File Generation & Email ---
if submitted and name and final_ward and user_email and confirm:
    df = pd.DataFrame(responses)
    base = "{}_{}".format(final_ward, today)
    paths = {
        "csv": SAVE_DIR / "{}_responses.csv".format(base),
        "excel": SAVE_DIR / "{}_responses.xlsx".format(base),
        "word": SAVE_DIR / "{}_responses.docx".format(base),
        "pdf": SAVE_DIR / "{}_responses.pdf".format(base),
        "zip": SAVE_DIR / "{}_hazard_survey.zip".format(base)
    }

    df.to_csv(paths["csv"], index=False)
    df.to_excel(paths["excel"], index=False)

    doc = Document()
    doc.add_heading("Hazard Risk Assessment for {}".format(final_ward), 0)
    for _, row in df.iterrows():
        doc.add_paragraph(
            "Hazard: {}\nLikelihood: {}\nImpact: {}\nDisruption: {}\nRisk Score: {}".format(
                row['Hazard'], row['Likelihood'], row['Impact'], row['Disruption'], row['Risk Score']
            )
        )
    doc.save(paths["word"])

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Hazard Risk Assessment for {}".format(final_ward), ln=True)
    for _, row in df.iterrows():
        text = "Hazard: {}\nLikelihood: {}\nImpact: {}\nDisruption: {}\nRisk Score: {}\n".format(
            row['Hazard'], row['Likelihood'], row['Impact'], row['Disruption'], row['Risk Score']
        )
        pdf.multi_cell(0, 10, txt=text)
    pdf.output(str(paths["pdf"]))

    with zipfile.ZipFile(paths["zip"], "w") as zipf:
        for file_key in ["csv", "excel", "word", "pdf"]:
            zipf.write(paths[file_key], arcname=paths[file_key].name)

    st.success("Survey submitted successfully! Files have been saved and emailed.")

# --- Disclaimer ---
st.markdown("---")
st.markdown("**Disclaimer:** The software is developed by Dingaan Mahlangu and should not be used without prior permission.")
